"""
LegalLens - Document Extractor
===============================
Extracts clean text from legal documents in PDF, DOCX, and TXT formats.
Includes security validation layer to detect and block:
    - File type spoofing (magic bytes validation)
    - Embedded JavaScript in PDFs
    - VBA macros in DOCX files
    - Embedded objects/files in PDFs
    - Prompt injection attempts targeting the LLM component

Supported formats:
    - PDF  (via PyMuPDF / fitz)
    - DOCX (via python-docx)
    - TXT  (plain text with encoding detection)
"""

import re
import zipfile
import logging
from pathlib import Path
from dataclasses import dataclass, field

import fitz  # PyMuPDF
import docx

logger = logging.getLogger("LegalLens.extractor")


# Security configuration


MAX_FILE_SIZE_MB = 50

# Magic bytes signatures for supported formats
MAGIC_BYTES = {
    ".pdf":  b"%PDF",
    ".docx": b"PK",      # DOCX is a ZIP archive
}

# Patterns indicating prompt injection attempts in extracted text
# These target common techniques used to manipulate LLM behavior
PROMPT_INJECTION_PATTERNS = [
    re.compile(r"ignore\s+(all\s+)?previous\s+instructions", re.IGNORECASE),
    re.compile(r"ignore\s+(all\s+)?above\s+instructions", re.IGNORECASE),
    re.compile(r"disregard\s+(all\s+)?previous", re.IGNORECASE),
    re.compile(r"you\s+are\s+now\s+(a|an)\s+", re.IGNORECASE),
    re.compile(r"system\s*:\s*override", re.IGNORECASE),
    re.compile(r"new\s+instructions?\s*:", re.IGNORECASE),
    re.compile(r"forget\s+(all\s+)?(your|previous)\s+", re.IGNORECASE),
    re.compile(r"always\s+(respond|answer|say|output|classify)\s+", re.IGNORECASE),
    re.compile(r"override\s+risk\s+(assessment|level|score)", re.IGNORECASE),
    re.compile(r"<\s*/?\s*system\s*>", re.IGNORECASE),
    re.compile(r"\[INST\]|\[/INST\]|\[SYSTEM\]", re.IGNORECASE),
    re.compile(r"###\s*(system|instruction|human|assistant)\s*:", re.IGNORECASE),
]


class SecurityError(Exception):
    """Raised when a security check fails during document processing."""
    pass


@dataclass
class ExtractedDocument:
    """Container for extracted document content."""
    filename: str
    format: str
    text: str
    page_count: int
    char_count: int
    word_count: int
    security_warnings: list = field(default_factory=list)


def extract(filepath: str) -> ExtractedDocument:
    """
    Extract text from a document file with security validation.

    Security checks (run BEFORE extraction):
        1. File size limit
        2. Magic bytes validation (is the file what it claims to be?)
        3. PDF: embedded JavaScript detection
        4. PDF: embedded objects/files detection
        5. DOCX: VBA macro detection

    Security checks (run AFTER extraction):
        6. Prompt injection pattern scanning

    Args:
        filepath: Path to PDF, DOCX, or TXT file.

    Returns:
        ExtractedDocument with cleaned text, metadata, and security warnings.

    Raises:
        SecurityError: If a blocking security issue is found (JS, macros, spoofed file).
        ValueError: If file format is not supported.
        FileNotFoundError: If file does not exist.
    """
    path = Path(filepath)
    warnings = []

    if not path.exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = path.suffix.lower()

    # --- Pre-extraction security checks ---
    _validate_file_size(path)
    _validate_magic_bytes(path, ext)

    if ext == ".pdf":
        _validate_pdf_security(path)
        text, pages = _extract_pdf(path)
    elif ext == ".docx":
        _validate_docx_security(path)
        text, pages = _extract_docx(path)
    elif ext == ".txt":
        text, pages = _extract_txt(path)
    else:
        raise ValueError(f"Unsupported format: {ext}. Use PDF, DOCX, or TXT.")

    text = _clean_text(text)

    # --- Post-extraction security check ---
    injection_warnings = _scan_prompt_injection(text)
    warnings.extend(injection_warnings)

    if warnings:
        for w in warnings:
            logger.warning("Security: %s in %s", w, path.name)

    return ExtractedDocument(
        filename=path.name,
        format=ext.lstrip("."),
        text=text,
        page_count=pages,
        char_count=len(text),
        word_count=len(text.split()),
        security_warnings=warnings,
    )


def _extract_pdf(path: Path) -> tuple:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(str(path))
    pages = []

    for page in doc:
        text = page.get_text("text")
        if text.strip():
            pages.append(text)

    doc.close()
    return "\n\n".join(pages), len(pages)


def _extract_docx(path: Path) -> tuple:
    """Extract text from DOCX using python-docx."""
    doc = docx.Document(str(path))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n\n".join(paragraphs), 1


def _extract_txt(path: Path) -> tuple:
    """Extract text from plain text file with encoding fallback."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            text = path.read_text(encoding=encoding)
            return text, 1
        except (UnicodeDecodeError, UnicodeError):
            continue

    # Last resort: read as bytes and decode with replacement
    raw = path.read_bytes()
    text = raw.decode("utf-8", errors="replace")
    return text, 1


def _clean_text(text: str) -> str:
    """
    Clean extracted text:
    - Normalize whitespace
    - Remove page number artifacts
    - Remove excessive blank lines
    - Fix common OCR/extraction artifacts
    """
    # Replace form feeds and vertical tabs
    text = text.replace("\f", "\n\n").replace("\v", "\n")

    # Remove standalone page numbers (e.g., "  12  " or "Page 12 of 45")
    text = re.sub(r"\n\s*Page\s+\d+\s*(of\s+\d+)?\s*\n", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"\n\s*-?\s*\d{1,3}\s*-?\s*\n", "\n", text)

    # Collapse multiple blank lines into two
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Normalize spaces (but keep newlines)
    text = re.sub(r"[^\S\n]+", " ", text)

    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)

    return text.strip()


# Security validation functions

def _validate_file_size(path: Path):
    """
    Block files exceeding the size limit.
    Prevents resource exhaustion from oversized uploads.
    """
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise SecurityError(
            f"File too large: {size_mb:.1f} MB (limit: {MAX_FILE_SIZE_MB} MB)"
        )


def _validate_magic_bytes(path: Path, ext: str):
    """
    Verify file content matches its extension using magic bytes.

    A file named 'contract.pdf' should start with '%PDF'.
    A file named 'contract.docx' should start with 'PK' (ZIP header).

    This catches file type spoofing — e.g., an EXE renamed to .pdf.
    TXT files have no fixed signature, so they are skipped.
    """
    expected = MAGIC_BYTES.get(ext)
    if expected is None:
        return  # No signature to check (e.g., .txt)

    with open(path, "rb") as f:
        header = f.read(max(len(expected), 8))

    if not header.startswith(expected):
        raise SecurityError(
            f"File type mismatch: {path.name} has extension '{ext}' "
            f"but content does not match expected format. "
            f"Expected header: {expected}, got: {header[:8]}"
        )


def _validate_pdf_security(path: Path):
    """
    Scan PDF for embedded JavaScript and embedded file objects.

    JavaScript in PDFs can:
        - Auto-open URLs (phishing, C2 communication)
        - Exploit PDF reader vulnerabilities
        - Execute arbitrary code in some viewers

    Embedded files in PDFs can:
        - Carry executables (EXE, DLL, SCR)
        - Contain secondary malicious documents
        - Bypass email/upload filters
    """
    doc = fitz.open(str(path))

    try:
        # Check 1: Embedded JavaScript
        for i in range(len(doc)):
            page = doc[i]
            # Check page-level JS actions
            for annot in page.annots() or []:
                info = annot.info
                if info and "javascript" in str(info).lower():
                    raise SecurityError(
                        f"Embedded JavaScript detected in {path.name} (page {i+1}). "
                        f"Blocked for security."
                    )

        # Check document-level JavaScript
        js = doc.get_page_labels()  # Access catalog
        catalog = doc.pdf_catalog()
        if catalog:
            catalog_str = str(doc.xref_object(catalog))
            if "/JavaScript" in catalog_str or "/JS " in catalog_str:
                raise SecurityError(
                    f"Document-level JavaScript detected in {path.name}. "
                    f"Blocked for security."
                )

        # Check 2: Embedded files
        embeds = doc.embfile_count()
        if embeds > 0:
            names = [doc.embfile_info(i).get("filename", "unknown") for i in range(embeds)]
            raise SecurityError(
                f"Embedded files detected in {path.name}: {names}. "
                f"Blocked for security."
            )

    finally:
        doc.close()


def _validate_docx_security(path: Path):
    """
    Scan DOCX for VBA macros.

    DOCX files are ZIP archives. If they contain 'vbaProject.bin',
    the document has embedded VBA macros. Legitimate contracts
    should never contain macros.

    Note: .docm is the macro-enabled format, but attackers rename
    .docm to .docx to bypass filters. The ZIP content reveals the truth.
    """
    try:
        with zipfile.ZipFile(str(path), "r") as z:
            names = z.namelist()

            # Check for VBA macro binary
            vba_files = [n for n in names if "vbaProject" in n or n.endswith(".bin")]
            if vba_files:
                raise SecurityError(
                    f"VBA macros detected in {path.name}: {vba_files}. "
                    f"Blocked for security."
                )

            # Check for suspicious embedded objects (OLE)
            ole_files = [n for n in names if "oleObject" in n or "activeX" in n]
            if ole_files:
                raise SecurityError(
                    f"Embedded OLE/ActiveX objects in {path.name}: {ole_files}. "
                    f"Blocked for security."
                )

    except zipfile.BadZipFile:
        raise SecurityError(
            f"Invalid DOCX structure: {path.name} is not a valid ZIP archive."
        )


def _scan_prompt_injection(text: str) -> list:
    """
    Scan extracted text for prompt injection patterns.

    Unlike other checks, this does NOT block the file — it returns
    warnings that are attached to the ExtractedDocument. The pipeline
    can then decide how to handle them (e.g., skip LLM analysis,
    flag in report, sanitize before sending to LLM).

    Why not block? Because some legitimate contracts might contain
    phrases like 'new instructions:' in a business context. We flag
    and let the pipeline decide.
    """
    warnings = []

    for pattern in PROMPT_INJECTION_PATTERNS:
        matches = pattern.findall(text)
        if matches:
            # Find the actual line containing the match for context
            for match in pattern.finditer(text):
                start = max(0, match.start() - 40)
                end = min(len(text), match.end() + 40)
                context = text[start:end].replace("\n", " ").strip()
                warnings.append(
                    f"Potential prompt injection: '...{context}...'"
                )
            break  # One warning per pattern type is enough

    return warnings
